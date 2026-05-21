import os
import re
import json
from docx import Document
from docx.oxml.ns import qn

def parse_math_element(element):
    text = ""
    for child in element:
        if child.tag.endswith('t'):
            text += child.text if child.text else ""
        elif child.tag.endswith('superscript') or child.tag.endswith('sup'):
            text += f"<sup>{parse_math_element(child)}</sup>"
        elif child.tag.endswith('subscript') or child.tag.endswith('sub'):
            text += f"<sub>{parse_math_element(child)}</sub>"
        elif len(child) > 0:
            text += parse_math_element(child)
    return text

def get_text_with_formatting(paragraph):
    p_html = ""
    for child in paragraph._p:
        if child.tag.endswith('r'):
            rPr = child.find(qn('w:rPr'))
            text_elem = child.find(qn('w:t'))
            if text_elem is not None and text_elem.text:
                text = text_elem.text
                if rPr is not None:
                    vertAlign = rPr.find(qn('w:vertAlign'))
                    if vertAlign is not None:
                        val = vertAlign.get(qn('w:val'))
                        if val == 'superscript':
                            p_html += f"<sup>{text}</sup>"
                            continue
                        elif val == 'subscript':
                            p_html += f"<sub>{text}</sub>"
                            continue
                p_html += text
        elif child.tag.endswith('oMath'):
            math_text = parse_math_element(child)
            if not math_text:
                math_text = "".join(child.itertext())
            math_text = math_text.replace("²", "<sup>2</sup>").replace("³", "<sup>3</sup>")
            p_html += f" <span class='math-expr'>{math_text}</span> "
    
    if not p_html.strip():
        p_html = paragraph.text if paragraph.text else ""
        
    p_html = re.sub(r'(\w+)/(\w+)', r'<span class="fraction"><span class="top">\1</span><span class="bottom">\2</span></span>', p_html)
    return p_html.strip()

def get_quizzes_by_letters(file_path, json_answers_path=None):
    try:
        doc = Document(file_path)
        paragraphs_html = []
        for p in doc.paragraphs:
            txt = get_text_with_formatting(p)
            if txt:
                paragraphs_html.append(txt)
                
        full_text = "\n".join(paragraphs_html)
        
        correct_answers = {}
        if json_answers_path and os.path.exists(json_answers_path):
            try:
                with open(json_answers_path, 'r', encoding='utf-8') as f:
                    answers_data = json.load(f)
                correct_answers = answers_data.get("fizika_answers", answers_data)
            except Exception as json_err:
                print(f"JSON yuklashda xatolik: {json_err}")

        question_blocks = re.split(r'(?=\b\d+\.)', full_text)
        quiz_data = []
        letter_to_idx = {"A": 0, "B": 1, "C": 2, "D": 3}

        for block in question_blocks:
            block = block.strip()
            if not block: 
                continue
            header_match = re.match(r'^(\d+)\.(.*)', block, re.DOTALL)
            if not header_match: 
                continue
                
            q_num = header_match.group(1).strip()
            rest_of_text = header_match.group(2).strip()
            
            idx_A = rest_of_text.find("A)")
            idx_B = rest_of_text.find("B)")
            idx_C = rest_of_text.find("C)")
            idx_D = rest_of_text.find("D)")
            
            if idx_A != -1 and idx_B != -1 and idx_C != -1 and idx_D != -1 and (idx_A < idx_B < idx_C < idx_D):
                question_content = rest_of_text[:idx_A].strip()
                opt_A = rest_of_text[idx_A + 2:idx_B].strip()
                opt_B = rest_of_text[idx_B + 2:idx_C].strip()
                opt_C = rest_of_text[idx_C + 2:idx_D].strip()
                opt_D = rest_of_text[idx_D + 2:].strip()
                
                options = [opt_A, opt_B, opt_C, opt_D]
                correct_letter = str(correct_answers.get(q_num, "A")).upper().strip()
                correct_index = letter_to_idx.get(correct_letter, 0)
                
                if question_content:
                    cleaned_question = re.sub(r'[ \t]+', ' ', question_content).strip()
                    cleaned_options = [re.sub(r'\|$', '', opt).strip() for opt in options]
                    
                    quiz_data.append({
                        "question": f"{q_num}. {cleaned_question}",
                        "options": cleaned_options,
                        "correct": correct_index
                    })
        return quiz_data
    except Exception as e:
        print(f"Fizika parser xatosi: {e}")
        return []

def get_quizzes(file_path):
    try:
        doc = Document(file_path)
        paragraphs = []
        for p in doc.paragraphs:
            t = get_text_with_formatting(p)
            if t: 
                paragraphs.append(t)
        full_text = "\n".join(paragraphs)
        raw_questions = full_text.split("++++")
        quiz_data = []

        for item in raw_questions:
            item = item.strip()
            if not item: 
                continue
            parts = item.split("====")
            if len(parts) < 2: 
                continue
            
            question_text = parts[0].strip()
            options_raw = [p.strip() for p in parts[1:] if p.strip()]
            
            correct_index = 0
            final_options = []
            for i, opt in enumerate(options_raw):
                clean_opt = opt.replace("#", "").strip()
                if opt.startswith("#"):
                    correct_index = i
                final_options.append(clean_opt)
            
            if len(final_options) >= 2:
                quiz_data.append({
                    "question": question_text,
                    "options": final_options,
                    "correct": correct_index
                })
        return quiz_data
    except Exception as e:
        print(f"Standart parser xatosi: {e}")
        return []

def get_quizzes_programming(file_path):
    try:
        doc = Document(file_path)
        quiz_data = []
        paragraphs = [get_text_with_formatting(p) for p in doc.paragraphs if p.text.strip()]
        
        i = 0
        while i < len(paragraphs):
            if "Savol" in paragraphs[i] or "савол" in paragraphs[i].lower():
                question_parts = []
                i += 1
                while i < len(paragraphs) and not paragraphs[i].startswith("—"):
                    question_parts.append(paragraphs[i])
                    i += 1
                
                question_text = " ".join(question_parts).strip()
                options = []
                while i < len(paragraphs) and paragraphs[i].startswith("—"):
                    option_text = paragraphs[i].replace("—", "").strip()
                    options.append(option_text)
                    i += 1
                
                if len(options) >= 2 and question_text:
                    quiz_data.append({
                        "question": question_text,
                        "options": options,
                        "correct": 0
                    })
            else:
                i += 1
        return quiz_data
    except Exception as e:
        print(f"Dasturlash parser xatosi: {e}")
        return []

def normalize_text(text):
    if not text: 
        return ""
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()

def calculate_similarity(text1, text2):
    norm1 = normalize_text(text1)
    norm2 = normalize_text(text2)
    if not norm1 or not norm2: 
        return 0
    if norm1 == norm2: 
        return 100
    words1 = set(norm1.split())
    words2 = set(norm2.split())
    union = len(words1.union(words2))
    if union == 0: 
        return 0
    return (len(words1.intersection(words2)) / union) * 100

def get_quizzes_english_pdf_docx(docx_path, pdf_path):
    try:
        import pdfplumber
        pdf_red_texts = []
        if os.path.exists(pdf_path):
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    chars = page.chars
                    current_word = []
                    for char in chars:
                        color = char.get("non_stroking_color") or char.get("stroking_color")
                        is_red = False
                        if color and len(color) == 3:
                            r, g, b = color
                            if (r > 0.6 and g < 0.4 and b < 0.4) or (isinstance(r, int) and r > 150 and g < 100 and b < 100):
                                is_red = True
                        if is_red:
                            current_word.append(char["text"])
                        else:
                            if current_word:
                                word_str = "".join(current_word).strip()
                                if word_str and not word_str.isdigit() and len(word_str) > 1:
                                    pdf_red_texts.append(word_str)
                                current_word = []
                    if current_word:
                        word_str = "".join(current_word).strip()
                        if word_str and not word_str.isdigit() and len(word_str) > 1:
                            pdf_red_texts.append(word_str)
        
        doc = Document(docx_path)
        clean_lines = []
        for p in doc.paragraphs:
            line = get_text_with_formatting(p)
            if line:
                if set(line).issubset({'_', '=', '+', '-', '*', ' ', '—'}) and len(line) > 1:
                    continue
                clean_lines.append(line)
        
        i = 0
        docx_quizzes = []
        while i < len(clean_lines):
            line = clean_lines[i]
            if any(phrase in line.lower() for phrase in ["choose the correct", "choose the word", "choose the best", "select the correct", "complete the sentence", "it is a group"]):
                question_text = line
                options = []
                i += 1
                while i < len(clean_lines):
                    current_line = clean_lines[i]
                    if any(phrase in current_line.lower() for phrase in ["choose the correct", "choose the word", "choose the best", "select the correct", "complete the sentence", "it is a group"]):
                        break
                    if current_line and len(current_line) < 200:
                        options.append(current_line)
                    i += 1
                if len(options) >= 2:
                    docx_quizzes.append({
                        "question": question_text,
                        "options": options
                    })
            else:
                i += 1
                
        quiz_data = []
        for docx_quiz in docx_quizzes:
            best_correct_idx = 0
            best_overall_score = 0
            for opt_idx, option in enumerate(docx_quiz["options"]):
                for pdf_text in pdf_red_texts:
                    similarity = calculate_similarity(pdf_text, option)
                    if similarity > best_overall_score:
                        best_overall_score = similarity
                        best_correct_idx = opt_idx
            
            quiz_data.append({
                "question": docx_quiz["question"],
                "options": docx_quiz["options"],
                "correct": best_correct_idx
            })
        return quiz_data
    except Exception as e:
        print(f"Ingliz tili PDF+DOCX parser xatosi: {e}")
        return []
